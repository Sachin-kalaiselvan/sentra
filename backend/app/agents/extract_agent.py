import os
import io
from app.graphs.state import SentraState

UPLOAD_DIR = os.getenv("UPLOAD_DIR", "/app/uploads")


def extract_agent(state: SentraState) -> SentraState:
    filename = state["filename"]
    file_path = os.path.join(UPLOAD_DIR, filename)
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    try:
        text = _parse_file(file_path, ext)
    except Exception as e:
        print(f"[Extract] Failed to parse {filename}: {e}")
        text = ""

    state["extracted_data"] = {
        "source": filename,
        "document_type": ext,
        "raw_text": text,
        "char_count": len(text),
    }
    return state


def _parse_file(path: str, ext: str) -> str:
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")

    if ext == "pdf":
        return _parse_pdf(path)
    elif ext == "docx":
        return _parse_docx(path)
    elif ext == "csv":
        return _parse_csv(path)
    elif ext in ("xlsx", "xls"):
        return _parse_excel(path)
    elif ext in ("txt", "md", "eml"):
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif ext in ("png", "jpg", "jpeg", "tiff"):
        return _parse_image(path)
    else:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()


def _parse_pdf(path: str) -> str:
    import pdfplumber
    texts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                texts.append(t)
            for table in page.extract_tables():
                if table:
                    header = table[0]
                    rows = table[1:]
                    line = " | ".join(str(c or "") for c in header)
                    texts.append(line)
                    for row in rows:
                        texts.append(" | ".join(str(c or "") for c in row))
    return "\n".join(texts)


def _parse_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _parse_csv(path: str) -> str:
    import pandas as pd
    df = pd.read_csv(path)
    return df.to_markdown(index=False)


def _parse_excel(path: str) -> str:
    import pandas as pd
    sheets = pd.read_excel(path, sheet_name=None)
    parts = []
    for name, df in sheets.items():
        parts.append(f"Sheet: {name}\n{df.to_markdown(index=False)}")
    return "\n\n".join(parts)


def _parse_image(path: str) -> str:
    import pytesseract
    from PIL import Image
    return pytesseract.image_to_string(Image.open(path))